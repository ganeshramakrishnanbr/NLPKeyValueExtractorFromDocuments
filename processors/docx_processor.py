"""
DOCX/DOC Document Processor using python-docx
"""

import asyncio
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
import chardet

class DOCXProcessor:
    """Process DOCX and DOC documents"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    async def process(self, file_path: Path) -> Dict:
        """Process DOCX/DOC document with text extraction and metadata"""
        
        try:
            # Extract text content
            text_content = await self._extract_text(file_path)
            
            # Extract metadata
            metadata = await self._extract_metadata(file_path)
            
            return {
                'text_content': text_content,
                'metadata': metadata,
                'extraction_method': 'docx_processor'
            }
            
        except Exception as e:
            self.logger.error(f"DOCX processing failed: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX document"""
        
        text_parts = []
        
        try:
            doc = Document(str(file_path))
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Text extraction failed: {e}")
            return ""
    
    async def _extract_metadata(self, file_path: Path) -> Dict:
        """Extract DOCX metadata"""
        
        metadata = {
            'file_size': file_path.stat().st_size,
            'file_type': 'docx' if file_path.suffix.lower() == '.docx' else 'doc'
        }
        
        try:
            doc = Document(str(file_path))
            core_props = doc.core_properties
            
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            })
            
        except Exception as e:
            self.logger.warning(f"Metadata extraction failed: {e}")
        
        return metadata