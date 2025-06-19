"""
Document Processing Module

This module provides comprehensive document processing capabilities for extracting 
text content from various file formats including PDF, DOCX, DOC, and Markdown files.
It also includes basic document classification functionality to categorize documents
based on their content.

Author: NLP Document Extraction Platform
Version: 2.0.0
"""

import pdfplumber
import docx
import docx2txt
from typing import Tuple
import logging

# Configure logging for document processing operations
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Advanced document processor for multi-format text extraction and classification.
    
    This class handles the extraction of text content from various document formats
    and provides intelligent document type classification based on content analysis.
    
    Supported Formats:
        - PDF (.pdf): Uses pdfplumber for robust text extraction
        - DOCX (.docx): Uses python-docx for modern Word documents  
        - DOC (.doc): Uses docx2txt for legacy Word documents
        - Markdown (.md): Direct text file processing
    
    Features:
        - Intelligent text extraction with format-specific optimizations
        - Document type classification using keyword analysis
        - Error handling for corrupted or unsupported files
        - Text preprocessing and normalization
    """
    
    def __init__(self):
        """
        Initialize the document processor with supported file formats.
        
        Sets up the processor with a list of supported file formats
        for validation and processing purposes.
        """
        self.supported_formats = ['pdf', 'docx', 'doc', 'md']
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from PDF, DOCX, DOC, or MD files
        
        Args:
            file_path: Path to the document file
            file_type: Type of file (pdf, docx, doc, md)
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If file type is not supported
            Exception: If text extraction fails
        """
        file_type = file_type.lower()
        
        if file_type not in self.supported_formats:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type == 'doc':
                return self._extract_doc_text(file_path)
            elif file_type == 'md':
                return self._extract_md_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_type} file: {str(e)}")
            raise Exception(f"Failed to extract text from {file_type} file: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        text_content = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            extracted_text = '\n\n'.join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n'.join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
        
    def _extract_doc_text(self, file_path: str) -> str:
        """Extract text from legacy DOC files using docx2txt"""
        try:
            extracted_text = docx2txt.process(file_path)
            if not extracted_text:
                raise Exception("No text could be extracted from DOC file")
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOC")
            return extracted_text
            
        except Exception as e:
            raise Exception(f"DOC processing error: {str(e)}")
    
    def _extract_md_text(self, file_path: str) -> str:
        """Extract text from Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            if not md_content.strip():
                raise Exception("No content found in Markdown file")
            
            # Remove Markdown syntax for cleaner text extraction
            import re
            
            # Remove headers (# ## ###)
            md_content = re.sub(r'^#{1,6}\s+', '', md_content, flags=re.MULTILINE)
            
            # Remove bold and italic markers
            md_content = re.sub(r'\*\*([^*]+)\*\*', r'\1', md_content)  # **bold**
            md_content = re.sub(r'\*([^*]+)\*', r'\1', md_content)      # *italic*
            md_content = re.sub(r'__([^_]+)__', r'\1', md_content)      # __bold__
            md_content = re.sub(r'_([^_]+)_', r'\1', md_content)        # _italic_
            
            # Remove code blocks
            md_content = re.sub(r'```[^`]*```', '', md_content, flags=re.DOTALL)
            md_content = re.sub(r'`([^`]+)`', r'\1', md_content)        # inline code
            
            # Remove links but keep link text
            md_content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', md_content)
            
            # Remove images
            md_content = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', md_content)
            
            # Remove horizontal rules
            md_content = re.sub(r'^[-*_]{3,}$', '', md_content, flags=re.MULTILINE)
            
            # Remove bullet points and numbers
            md_content = re.sub(r'^\s*[-*+]\s+', '', md_content, flags=re.MULTILINE)
            md_content = re.sub(r'^\s*\d+\.\s+', '', md_content, flags=re.MULTILINE)
            
            # Clean up extra whitespace
            md_content = re.sub(r'\n\s*\n', '\n\n', md_content)
            md_content = md_content.strip()
            
            logger.info(f"Successfully extracted {len(md_content)} characters from Markdown")
            return md_content
            
        except Exception as e:
            raise Exception(f"Markdown processing error: {str(e)}")
    
    def classify_document_type(self, text: str) -> str:
        """
        Basic document type classification using keyword matching
        
        Args:
            text: Extracted text from document
            
        Returns:
            Document type classification
        """
        # Document type keywords
        keywords = {
            "employment_document": [
                "employment", "employee", "employment contract", "job offer", "work agreement",
                "salary", "wage", "position", "department", "employer", "hire", "staff"
            ],
            "financial_document": [
                "bank statement", "financial", "account", "balance", "transaction",
                "deposit", "withdrawal", "investment", "loan", "credit", "debit"
            ],
            "legal_document": [
                "contract", "agreement", "legal", "terms", "conditions", "clause",
                "party", "witness", "notary", "court", "jurisdiction"
            ],
            "medical_document": [
                "medical", "health", "patient", "doctor", "physician", "treatment",
                "diagnosis", "prescription", "hospital", "clinic", "healthcare"
            ],
            "insurance_document": [
                "insurance", "policy", "coverage", "premium", "deductible", "claim",
                "beneficiary", "insured", "insurer", "protection"
            ],
            "educational_document": [
                "education", "school", "university", "student", "course", "grade",
                "transcript", "diploma", "certificate", "academic"
            ],
            "identification_document": [
                "identification", "id", "passport", "license", "permit", "certificate",
                "registration", "official", "government", "issued"
            ]
        }
        
        text_lower = text.lower()
        
        # Count keyword matches for each document type
        type_scores = {}
        for doc_type, keyword_list in keywords.items():
            score = sum(1 for keyword in keyword_list if keyword in text_lower)
            if score > 0:
                type_scores[doc_type] = score
        
        # Return the document type with the highest score
        if type_scores:
            best_match = max(type_scores.keys(), key=lambda k: type_scores[k])
            logger.info(f"Document classified as: {best_match} (score: {type_scores[best_match]})")
            return best_match
        
        # Default to unknown if no keywords match
        logger.info("Document type could not be determined - classified as unknown")
        return "unknown"
